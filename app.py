from flask import Flask, render_template, request, jsonify, send_file
import os
from docx import Document  # <-- This line was missing

# Import the new test case generation function
from test_case_generation import generate_test_cases
# --- NEW: Import from our new document parser file ---
from document_parser import read_file_content, parse_document_for_tests
from test_executor import run_tests


app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/pipeline')
def pipeline():
    return render_template('pipeline.html')

# Product pages
@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/integrations')
def integrations():
    return render_template('integrations.html')

# Company pages
@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/careers')
def careers():
    return render_template('careers.html')

@app.route('/blog')
def blog():
    return render_template('blog.html')

@app.route('/support')
def support():
    return render_template('support.html')

# Legal pages
@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/terms')
def terms():
    return render_template('terms.html')

@app.route('/security')
def security():
    return render_template('security.html')

@app.route('/compliance')
def compliance():
    return render_template('compliance.html')

@app.route('/cookies')
def cookies():
    return render_template('cookies.html')

@app.route('/api/generate-test', methods=['POST'])
def handle_generate_test():
    """
    API endpoint for generating test cases using the LLM.
    """
    try:
        data = request.get_json()
        test_type = data.get('test_type')
        
        if not test_type:
            return jsonify({'status': 'error', 'message': 'test_type is required'}), 400

        test_cases = generate_test_cases(test_type, data)
        
        if isinstance(test_cases, dict) and 'error' in test_cases:
            return jsonify({'status': 'error', 'message': test_cases['details']}), 500

        return jsonify({
            'status': 'success',
            'message': 'Test cases generated successfully!',
            'tests': test_cases
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/parse-tests-from-file', methods=['POST'])
def parse_tests_from_file_endpoint():
    """
    API endpoint for parsing test cases from an uploaded DOCX or PDF file.
    """
    try:
        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': 'No file part in the request'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': 'No file selected'}), 400
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)
        if file_size > app.config['MAX_CONTENT_LENGTH']:
            return jsonify({'status': 'error', 'message': f"File exceeds the {app.config['MAX_CONTENT_LENGTH'] / 1024 / 1024}MB size limit."}), 413
        file_content = read_file_content(file)
        if not file_content.strip():
            return jsonify({'status': 'error', 'message': 'Could not extract any text from the uploaded file or the file type is unsupported.'}), 400
        all_tests = parse_document_for_tests(file_content)
        if not all_tests:
            return jsonify({'status': 'error', 'message': 'No valid test cases could be extracted from the document.'}), 500
        return jsonify({
            'status': 'success',
            'message': 'Test cases parsed successfully!',
            'tests': all_tests
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': f"An unexpected server error occurred: {str(e)}"}), 500


@app.route('/api/run-test', methods=['POST'])
def run_test():
    """API endpoint for running test cases"""
    try:
        data = request.get_json()
        website_url = data.get('website_url')
        test_cases = data.get('test_cases', [])
        if not website_url:
            return jsonify({'status': 'error', 'message': 'website_url is required'}), 400
        if not isinstance(test_cases, list) or len(test_cases) == 0:
            return jsonify({'status': 'error', 'message': 'test_cases must be a non-empty array'}), 400

        results = run_tests(website_url, test_cases)
        return jsonify({
            'status': 'success',
            'message': f'Tests executed on website: {website_url}',
            'results': results
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """API endpoint for file uploads for the generation flow (in-memory only)"""
    try:
        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': 'No file uploaded'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': 'No file selected'}), 400
        # No disk write, just acknowledge upload
        return jsonify({
            'status': 'success',
            'message': 'File received in memory',
            'filename': file.filename
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/download-tests', methods=['POST'])
def download_tests():
    """
    API endpoint for generating and downloading a DOCX of test cases (in-memory, no disk write).
    """
    from io import BytesIO
    try:
        data = request.get_json()
        test_cases = data.get('test_cases', [])
        if not test_cases:
            return jsonify({'status': 'error', 'message': 'No test cases provided'}), 400
        document = Document()
        document.add_heading('Generated Test Cases', 0)
        for test in test_cases:
            document.add_heading(f"ID: {test.get('id', 'N/A')} - {test.get('name', 'No Name')}", level=1)
            document.add_paragraph(f"Description: {test.get('description', 'No Description')}")
            document.add_paragraph(f"Type: {test.get('type', 'N/A')}")
            document.add_paragraph(f"Selector: {test.get('selector', 'N/A')}")
            document.add_paragraph()
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, download_name="generated_test_cases.docx")
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/api/download-results', methods=['POST'])
def download_results():
    """
    API endpoint for generating and downloading a DOCX of test execution results (in-memory, no disk write).
    """
    from io import BytesIO
    try:
        data = request.get_json()
        test_results = data.get('test_results', [])
        if not test_results:
            return jsonify({'status': 'error', 'message': 'No test results provided'}), 400
        document = Document()
        document.add_heading('Test Execution Results', 0)
        for result in test_results:
            name = result.get('name', 'Unnamed Test')
            status = result.get('status', 'unknown').upper()
            message = result.get('message', '')
            document.add_heading(name, level=1)
            document.add_paragraph(f"Status: {status}")
            if message:
                document.add_paragraph(f"Details: {message}")
            document.add_paragraph()
        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        return send_file(file_stream, as_attachment=True, download_name="test_execution_results.docx")
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/submit-contact', methods=['POST'])
def submit_contact():
    """
    Endpoint to handle contact form submissions.
    """
    try:
        # Get data from the form
        data = request.form
        name = data.get('name')
        email = data.get('email')
        message = data.get('message')

        if not name or not email or not message:
            return jsonify({'status': 'error', 'message': 'All fields are required.'}), 400

        # Placeholder for saving data logic
        # You can implement your own logic here

        return jsonify({'status': 'success', 'message': 'We will contact you soon!'}), 200

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
